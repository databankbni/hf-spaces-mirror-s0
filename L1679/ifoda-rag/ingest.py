"""
Ingestion pipeline: extracts documents, chunks, embeds, stores in ChromaDB.
"""
import os, re, hashlib
from typing import List, Dict
from dataclasses import dataclass, field

import chromadb
from chromadb.utils import embedding_functions
import fitz
from docx import Document as DocxDoc
import openpyxl

from config import *

@dataclass
class Chunk:
    id: str
    text: str
    metadata: Dict = field(default_factory=dict)

class IFODAIngestor:
    def __init__(self):
        print(f"[INGEST] Loading embedding model: {EMBEDDING_MODEL}")
        self.embedding_fn = embedding_functions.SentenceTransformerEmbeddingFunction(
            model_name=EMBEDDING_MODEL, device="cpu")
        self.client = chromadb.PersistentClient(path=VECTOR_DB_DIR)
        self._init_collection()
        self.chunks: List[Chunk] = []
        self._counter = 0

    def _uid(self, prefix: str) -> str:
        self._counter += 1
        return f"{prefix}_{self._counter}"

    def _init_collection(self):
        try: self.client.delete_collection(CHROMA_COLLECTION)
        except: pass
        self.collection = self.client.create_collection(
            name=CHROMA_COLLECTION,
            embedding_function=self.embedding_fn,
            metadata={"hnsw:space": "cosine"})
        print(f"[INGEST] Created collection: {CHROMA_COLLECTION}")

    def ingest_excel(self, filepath: str):
        print(f"[INGEST] Excel: {filepath}")
        wb = openpyxl.load_workbook(filepath)
        ws = wb["Лист1"]
        cnt = 0
        for row in ws.iter_rows(values_only=True):
            if not row[1] or row[1] == "Код": continue
            p = {"code": str(row[1]).strip(), "name_logix": str(row[2] or "").strip(),
                 "name_nsi": str(row[3] or "").strip(), "unit": str(row[4] or "").strip(),
                 "category_1": str(row[5] or "").strip(), "category_2": str(row[6] or "").strip(),
                 "category_3": str(row[7] or "").strip()}
            text = f"SKU: {p['code']}\nProduct: {p['name_logix']} / {p['name_nsi']}\nUnit: {p['unit']}\nCategory: {p['category_1']} → {p['category_2']} → {p['category_3']}"
            self.chunks.append(Chunk(id=self._uid("sku"), text=text, metadata={
                "source": os.path.basename(filepath), "source_type": "excel_sku",
                "product_code": p["code"], "product_name": p["name_logix"],
                "product_name_ru": p["name_nsi"], "unit": p["unit"],
                "category_1": p["category_1"], "category_2": p["category_2"],
                "category_3": p["category_3"], "doc_type": "structured_product"}))
            cnt += 1
        print(f"[INGEST] Excel: {cnt} products")

    def ingest_docx(self, filepath: str):
        print(f"[INGEST] DOCX: {filepath}")
        doc = DocxDoc(filepath)
        product_type_re = re.compile(r'(Insecticide|Fungicide|Herbicide|Fertilizer|Defoliant|Nematicide|Bactericide|Seed protectant|Plant growth stimulator|Adjuvant|Surfactant)', re.I)
        paragraphs = [p.text.strip() for p in doc.paragraphs if p.text.strip()]
        entries = []; current_entry = []; current_name = None
        for i, text in enumerate(paragraphs):
            is_header = False
            if len(text) < 60 and not text.startswith(('Culture','Hazard','Dosage','Winter','Cotton','Apple','Tomato','Potato','Grape','Wheat','Corn','Rice','Onion','Peach','Mulberry','Spraying','Etching','Last')):
                if i+1 < len(paragraphs) and product_type_re.match(paragraphs[i+1]):
                    is_header = True
            if is_header and current_entry:
                entries.append({"name": current_name or "Unknown", "text": "\n".join(current_entry)})
                current_entry = []; current_name = text; current_entry.append(text)
            else:
                if not current_entry and not is_header: current_name = text
                current_entry.append(text)
        if current_entry: entries.append({"name": current_name or "Unknown", "text": "\n".join(current_entry)})
        for entry in entries:
            if len(entry["text"]) < 30: continue
            self.chunks.append(Chunk(
                id=self._uid("docx"), text=entry["text"], metadata={"source": os.path.basename(filepath),
                "source_type": "docx_catalog", "product_name": entry["name"], "doc_type": "product_description"}))
        for ti, table in enumerate(doc.tables):
            rows = [" | ".join(c.text.strip() for c in row.cells) for row in table.rows]
            ct = "\n".join(rows)
            if len(ct) > 20:
                self.chunks.append(Chunk(id=self._uid("docx_table"), text=f"DOSAGE TABLE:\n{ct}",
                    metadata={"source": os.path.basename(filepath), "source_type": "docx_table",
                    "doc_type": "dosage_table", "table_index": ti}))
        print(f"[INGEST] DOCX: {len(entries)} product entries + {len(doc.tables)} tables")

    def ingest_pdf(self, filepath: str, source_type: str = "pdf"):
        print(f"[INGEST] PDF: {filepath}")
        doc = fitz.open(filepath)
        full_text = ""
        for pg in range(doc.page_count):
            t = doc[pg].get_text("text")
            if t.strip(): full_text += t + "\n"
        doc.close()
        if not full_text.strip(): return
        lang = "ru" if any(c in full_text[:500] for c in "агдеёжзийклмнопрстуфхцчшщъыьэюя") else "en"
        lines = full_text.split("\n")
        chunks_prod = []; curr = []; curr_prod = "General"
        for line in lines:
            s = line.strip()
            if not s: continue
            is_prod = (len(s) < 80 and (s.isupper() or re.match(r'^[A-ZА-Я][A-Za-zА-Яа-я\- ]{2,60}$', s)) and
                       not s.startswith(('Для','При','Цель','Культура','Culture','Hazard','Dosage','Application',
                                         'Namangan','Tel','@','www','What','When','Органо')))
            if is_prod and len(curr) > 5:
                chunks_prod.append((curr_prod, "\n".join(curr)))
                curr = [s]; curr_prod = s
            else: curr.append(s)
        if curr: chunks_prod.append((curr_prod, "\n".join(curr)))
        for pn, txt in chunks_prod:
            if len(txt) < 100: continue
            self.chunks.append(Chunk(
                id=self._uid("pdf"), text=txt,
                metadata={"source": os.path.basename(filepath), "source_type": source_type,
                "product_name": pn, "doc_type": "product_catalog", "language": lang}))
        print(f"[INGEST] PDF: {len(chunks_prod)} chunks")

    def embed_and_store(self, batch_size=100):
        print(f"[INGEST] Embedding {len(self.chunks)} chunks...")
        for i in range(0, len(self.chunks), batch_size):
            batch = self.chunks[i:i+batch_size]
            self.collection.add(ids=[c.id for c in batch], documents=[c.text for c in batch],
                               metadatas=[c.metadata for c in batch])
            print(f"[INGEST] Batch {i//batch_size+1}: {len(batch)} chunks")
        print(f"[INGEST] Total: {self.collection.count()} documents in ChromaDB")

    def run_full_ingestion(self):
        print("="*60); print("IFODA RAG — INGESTION"); print("="*60)
        for fname in ["Б-2024 - названия_товаров.xlsx", "КАТАЛОГ ПРЕПАРАТОВ (2)(en) (2).docx"]:
            fp = os.path.join(DATA_DIR, fname)
            if os.path.exists(fp):
                if fname.endswith('.xlsx'): self.ingest_excel(fp)
                else: self.ingest_docx(fp)
        for fname, st in [("[RE]RU УДАБРЕНИЯ (2).pdf","pdf_fert_ru"),("[RE]RU УДАБРЕНИЯ (3).pdf","pdf_fert_ru"),
                          ("EN.pdf","pdf_fert_en"),("ifoda EN.pdf","pdf_cat_en"),
                          ("[RE-1]_Full_pages_Product-Catalog.pdf","pdf_registry")]:
            fp = os.path.join(DATA_DIR, fname)
            if os.path.exists(fp): self.ingest_pdf(fp, st)
        self.embed_and_store()
        print("="*60); print(f"[INGEST] DONE. {self.collection.count()} docs in {CHROMA_COLLECTION}"); print("="*60)

if __name__ == "__main__":
    IFODAIngestor().run_full_ingestion()
