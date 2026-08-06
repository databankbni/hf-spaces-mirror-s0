from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches, Pt


SOURCE = Path(__file__).parents[1] / "FCC_diesel_March_2025_answer.md"
TARGET = Path(__file__).parents[1] / "FCC_diesel_March_2025_answer.docx"


def add_inline(paragraph, text):
    parts = text.split("**")
    for index, part in enumerate(parts):
        run = paragraph.add_run(part.replace("`", ""))
        run.bold = index % 2 == 1


document = Document()
section = document.sections[0]
section.top_margin = Inches(0.7)
section.bottom_margin = Inches(0.7)
section.left_margin = Inches(0.75)
section.right_margin = Inches(0.75)

styles = document.styles
styles["Normal"].font.name = "Aptos"
styles["Normal"].font.size = Pt(10)
for name, size in (("Title", 20), ("Heading 1", 15), ("Heading 2", 12), ("Heading 3", 11)):
    styles[name].font.name = "Aptos Display"
    styles[name].font.size = Pt(size)

lines = SOURCE.read_text(encoding="utf-8").splitlines()
i = 0
while i < len(lines):
    line = lines[i].strip()
    if not line:
        i += 1
        continue

    if line.startswith("|") and i + 1 < len(lines) and lines[i + 1].strip().startswith("|---"):
        rows = []
        while i < len(lines) and lines[i].strip().startswith("|"):
            rows.append([cell.strip() for cell in lines[i].strip().strip("|").split("|")])
            i += 1
        rows.pop(1)
        table = document.add_table(rows=len(rows), cols=len(rows[0]))
        table.style = "Table Grid"
        for row_index, row in enumerate(rows):
            for col_index, value in enumerate(row):
                cell = table.cell(row_index, col_index)
                cell.text = value.replace("**", "").replace("`", "")
                for run in cell.paragraphs[0].runs:
                    run.bold = row_index == 0
                    run.font.size = Pt(8.5)
        continue

    if line.startswith("# "):
        paragraph = document.add_paragraph(style="Title")
        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        add_inline(paragraph, line[2:])
    elif line.startswith("## "):
        add_inline(document.add_paragraph(style="Heading 1"), line[3:])
    elif line.startswith("### "):
        add_inline(document.add_paragraph(style="Heading 2"), line[4:])
    elif line.startswith("- "):
        add_inline(document.add_paragraph(style="List Bullet"), line[2:])
    elif line.startswith("> "):
        paragraph = document.add_paragraph()
        paragraph.paragraph_format.left_indent = Inches(0.35)
        paragraph.paragraph_format.right_indent = Inches(0.35)
        add_inline(paragraph, line[2:])
        for run in paragraph.runs:
            run.italic = True
    else:
        add_inline(document.add_paragraph(), line)
    i += 1

document.core_properties.title = "FCC Fractionator Diesel Production — March 2025"
document.core_properties.subject = "ArcticDB-supported operating assessment"
document.save(TARGET)
print(TARGET)
