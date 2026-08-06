from flask import Flask, request, jsonify, Response
import sqlite3
import os
import re
import csv
from io import StringIO
from collections import Counter
from dotenv import load_dotenv

load_dotenv()
try:
    from huggingface_hub import hf_hub_download
except ImportError:
    hf_hub_download = None

DB_FILE = 'truku_dict.db'

def get_corpus_col(cursor):
    try:
        cols = [row[1] for row in cursor.execute("PRAGMA table_info(dictionary)").fetchall()]
        for col_name in ['corpus_source', 'source', 'corpus', 'category', 'type']:
            if col_name in cols:
                return col_name
    except Exception:
        pass
    return None

def has_column(cursor, table, column):
    try:
        cols = [row[1] for row in cursor.execute(f"PRAGMA table_info({table})").fetchall()]
        return column in cols
    except Exception:
        return False

# Download private DB from Hugging Face if configured
if os.environ.get('HF_TOKEN') and os.environ.get('DB_REPO_ID') and hf_hub_download:
    print("Downloading private database from Hugging Face dataset...")
    try:
        hf_hub_download(
            repo_id=os.environ.get('DB_REPO_ID'),
            filename=DB_FILE,
            repo_type="dataset",
            local_dir=".",
            token=os.environ.get('HF_TOKEN')
        )
        print("Database downloaded securely.")
    except Exception as e:
        print(f"Failed to download database: {e}")

app = Flask(__name__, static_folder='static')

stats_cache = None

all_source_words = set()
root_family_map = {}
word_to_root = {}
morphology_built = False

def extract_root(word, valid_words):
    if not word or len(word) < 3: return word
    w = word.lower()
    possible_roots = {w}
    
    prefixes = ['m', 'p', 'pn', 'mn', 's', 'sn', 't', 'tn', 'en', 'em', 'emp']
    suffixes = ['an', 'un']
    
    # Prefix
    for p in prefixes:
        if w.startswith(p):
            stem = w[len(p):]
            if stem in valid_words: possible_roots.add(stem)
    
    # Suffix
    for s in suffixes:
        if w.endswith(s):
            stem = w[:-len(s)]
            if stem in valid_words: possible_roots.add(stem)
            
    # Prefix + Suffix
    for p in prefixes:
        for s in suffixes:
            if w.startswith(p) and w.endswith(s) and len(w) > len(p) + len(s):
                stem = w[len(p):-len(s)]
                if stem in valid_words: possible_roots.add(stem)
                
    # Infix (usually n or m after first consonant)
    if len(w) > 3 and w[0] in 'ptks' and w[1] in 'mn':
        stem1 = w[0] + w[2:]
        if stem1 in valid_words: possible_roots.add(stem1)
        stem2 = w[2:]
        if stem2 in valid_words: possible_roots.add(stem2)
        
    return sorted(list(possible_roots), key=len)[0]

def build_morphology_cache():
    global all_source_words, root_family_map, word_to_root, morphology_built
    if morphology_built: return
    
    conn = sqlite3.connect(DB_FILE)
    cursor = conn.cursor()
    cursor.execute("SELECT DISTINCT source_word FROM dictionary WHERE source_word IS NOT NULL AND source_word != ''")
    rows = cursor.fetchall()
    conn.close()
    
    all_source_words = {re.sub(r'[^\w\-]', '', r[0]).lower() for r in rows}
    all_source_words.discard('')
    
    for w in all_source_words:
        word_to_root[w] = extract_root(w, all_source_words)
        
    for w, root in word_to_root.items():
        if root not in root_family_map:
            root_family_map[root] = set()
        root_family_map[root].add(w)
        
    morphology_built = True

def get_morphology_for_word(word):
    build_morphology_cache()
    clean = re.sub(r'[^\w\-]', '', word).lower()
    if not clean or clean not in word_to_root: return None
    root = word_to_root[clean]
    family = sorted(list(root_family_map.get(root, set())))
    if len(family) <= 1: return None
    return {
        "root": root,
        "family": family
    }

class SafeRow:
    """sqlite3.Row wrapper：當欄位不存在時自動回傳預設值，防止 IndexError。"""
    def __init__(self, row, col_map):
        self._row = row
        self._col_map = col_map  # 實際存在的欄位名稱 set
    def __getitem__(self, key):
        if isinstance(key, (int, slice)):
            return self._row[key]
        if key in self._col_map:
            return self._row[key]
        # 常見別名對應
        aliases = {
            'corpus_source': ['source', 'corpus', 'category', 'type'],
            'detailed_source': ['detailed_source', 'subcategory'],
        }
        for alt in aliases.get(key, []):
            if alt in self._col_map:
                return self._row[alt]
        return ''  # 欄位不存在時回傳空字串
    def get(self, key, default=''):
        try:
            return self[key] or default
        except Exception:
            return default
    def keys(self):
        return self._row.keys()

def smart_row_factory(cursor, row):
    col_names = {d[0] for d in cursor.description}
    base = sqlite3.Row(cursor, row)
    return SafeRow(base, col_names)

def get_db_connection():
    conn = sqlite3.connect(DB_FILE)
    conn.row_factory = smart_row_factory
    return conn

def tokenize(text):
    if not text: return []
    return [re.sub(r'[^\w\-]', '', w).lower() for w in text.split() if re.sub(r'[^\w\-]', '', w).lower()]

def get_kwic_regex(sentence, regex, is_exact=True):
    if not sentence or not regex: return sentence
    tokens = sentence.split()
    target_idx = -1
    for i, token in enumerate(tokens):
        clean_token = re.sub(r'[^\w\-]', '', token)
        if regex.search(clean_token if is_exact else token):
            target_idx = i
            break
            
    if target_idx == -1: return sentence
    start_idx = max(0, target_idx - 3)
    end_idx = min(len(tokens), target_idx + 4)
    result = " ".join(tokens[start_idx:end_idx])
    if start_idx > 0: result = "... " + result
    if end_idx < len(tokens): result = result + " ..."
    return result

def get_kwic_proximity(sentence, word1, word2):
    if not sentence: return sentence
    tokens = sentence.split()
    w1_lower = word1.lower()
    w2_lower = word2.lower()
    
    target_idx = -1
    for i, token in enumerate(tokens):
        clean = re.sub(r'[^\w\-]', '', token).lower()
        if clean == w1_lower or clean == w2_lower:
            target_idx = i
            break
            
    if target_idx == -1: return sentence
    start_idx = max(0, target_idx - 4)
    end_idx = min(len(tokens), target_idx + 5)
    result = " ".join(tokens[start_idx:end_idx])
    if start_idx > 0: result = "... " + result
    if end_idx < len(tokens): result = result + " ..."
    return result

def check_proximity(sentence, word1, word2, distance):
    tokens = tokenize(sentence)
    w1_lower = word1.lower()
    w2_lower = word2.lower()
    idx1 = [i for i, t in enumerate(tokens) if t == w1_lower]
    idx2 = [i for i, t in enumerate(tokens) if t == w2_lower]
    for i in idx1:
        for j in idx2:
            if abs(i - j) - 1 <= distance:
                return True
    return False

def get_collocations_regex(sentences, regex, top_n=5):
    colloc_counter = Counter()
    for sent in sentences:
        tokens = set(tokenize(sent))
        matched_tokens = {t for t in tokens if regex.search(t)}
        for t in tokens:
            if t and t not in matched_tokens and len(t) >= 2:
                colloc_counter[t] += 1
    return [word for word, count in colloc_counter.most_common(top_n)]

def get_collocations_prox(sentences, word1, word2, top_n=5):
    colloc_counter = Counter()
    w1 = word1.lower()
    w2 = word2.lower()
    for sent in sentences:
        tokens = set(tokenize(sent))
        for t in tokens:
            if t and t != w1 and t != w2 and len(t) >= 2:
                colloc_counter[t] += 1
    return [word for word, count in colloc_counter.most_common(top_n)]

def get_ngrams_regex(sentences, regex, n=2, top_n=3):
    ngram_counter = Counter()
    for sent in sentences:
        tokens = tokenize(sent)
        for i, t in enumerate(tokens):
            if regex.search(t):
                start_range = max(0, i - n + 1)
                end_range = min(i, len(tokens) - n)
                for start_idx in range(start_range, end_range + 1):
                    ngram = tuple(tokens[start_idx : start_idx + n])
                    if len(ngram) == n:
                        ngram_counter[ngram] += 1
    return [{"ngram": " ".join(ngram), "count": count} for ngram, count in ngram_counter.most_common(top_n)]

def get_position_stats_regex(sentences, regex):
    counts = {'initial': 0, 'medial': 0, 'final': 0}
    for sentence in sentences:
        matches = list(regex.finditer(sentence))
        for match in matches:
            start = match.start()
            end = match.end()
            prefix = sentence[:start].rstrip()
            suffix = sentence[end:].lstrip()
            is_initial = (not prefix) or (prefix[-1] in ',.?!;:')
            is_final = (not suffix) or (suffix[0] in ',.?!;:')
            if is_initial: counts['initial'] += 1
            elif is_final: counts['final'] += 1
            else: counts['medial'] += 1
    return counts

def perform_search(query, search_type, match_mode, limit=1000, corpora=None):
    if not query:
        return {"results": [], "total_occurrences": 0, "exact_occurrences": 0, "partial_occurrences": 0, "collocations": [], "ngrams": {}, "positions": {}, "morphology": None}
        
    is_proximity = False
    word1, word2, distance = None, None, 0
    regex_exact = None
    regex_partial = None
    
    corpus_condition = ""
    params = []
    if corpora:
        conn_temp = get_db_connection()
        c_col = get_corpus_col(conn_temp.cursor())
        conn_temp.close()
        if c_col:
            placeholders = ','.join(['?'] * len(corpora))
            corpus_condition = f" AND {c_col} IN ({placeholders})"
            params = list(corpora)
    
    if ' NEAR:' in query:
        is_proximity = True
        parts = query.split(' NEAR:')
        word1 = parts[0].strip()
        n_str, word2 = parts[1].strip().split(' ', 1)
        distance = int(n_str)
        word2 = word2.strip()
    else:
        if '*' in query:
            safe_q = re.escape(query).replace(r'\*', r'\w*')
        else:
            safe_q = re.escape(query)
            
        if search_type in ['truku', 'source']:
            regex_exact = re.compile(r'\b' + safe_q + r'\b', re.IGNORECASE)
            regex_partial = re.compile(safe_q, re.IGNORECASE)
        else:
            regex_exact = re.compile(safe_q, re.IGNORECASE)
            regex_partial = re.compile(safe_q, re.IGNORECASE)
            
    conn = get_db_connection()
    cursor = conn.cursor()
    
    if is_proximity:
        sql = f"SELECT * FROM dictionary WHERE truku_sentence LIKE ? AND truku_sentence LIKE ?{corpus_condition}"
        all_rows = cursor.execute(sql, (f"%{word1}%", f"%{word2}%") + tuple(params)).fetchall()
    else:
        sql_like = query.replace('*', '%')
        if search_type == 'truku': sql = f"SELECT * FROM dictionary WHERE truku_sentence LIKE ?{corpus_condition}"
        elif search_type == 'chinese': sql = f"SELECT * FROM dictionary WHERE chinese_translation LIKE ?{corpus_condition}"
        else: sql = f"SELECT * FROM dictionary WHERE source_word LIKE ?{corpus_condition}"
        all_rows = cursor.execute(sql, (f"%{sql_like}%",) + tuple(params)).fetchall()
        
    conn.close()
    
    exact_rows = []
    partial_rows = []
    
    for row in all_rows:
        if is_proximity:
            if search_type == 'truku' and check_proximity(row['truku_sentence'], word1, word2, distance):
                exact_rows.append(row)
        else:
            text = row['truku_sentence'] if search_type == 'truku' else (row['chinese_translation'] if search_type == 'chinese' else row['source_word'])
            if text:
                if regex_exact.search(text): exact_rows.append(row)
                elif match_mode == 'partial' and regex_partial.search(text): partial_rows.append(row)
                elif match_mode == 'partial' and '*' in query: partial_rows.append(row)
                
    final_rows = exact_rows if match_mode == 'exact' else exact_rows + partial_rows
    
    results = []
    iter_rows = final_rows if limit is None else final_rows[:limit]
    
    for row in iter_rows:
        item = dict(row)
        text = item['truku_sentence'] if search_type == 'truku' else (item['chinese_translation'] if search_type == 'chinese' else item['source_word'])
        if is_proximity:
            item['is_exact'] = True
            if search_type == 'truku':
                item['kwic_sentence'] = get_kwic_proximity(item['truku_sentence'], word1, word2)
                item['positions'] = []
            else:
                item['kwic_sentence'] = item['truku_sentence']
                item['positions'] = []
        else:
            is_ex = bool(text and regex_exact.search(text))
            item['is_exact'] = is_ex
            if search_type == 'truku':
                item['kwic_sentence'] = get_kwic_regex(item['truku_sentence'], regex_exact if is_ex else regex_partial, is_exact=is_ex)
                item['positions'] = []
                if is_ex:
                    matches = list(regex_exact.finditer(item['truku_sentence']))
                    p_list = []
                    for m in matches:
                        p = item['truku_sentence'][:m.start()].rstrip()
                        s = item['truku_sentence'][m.end():].lstrip()
                        if not p or p[-1] in ',.?!;:': p_list.append('initial')
                        elif not s or s[0] in ',.?!;:': p_list.append('final')
                        else: p_list.append('medial')
                    item['positions'] = list(set(p_list))
            else:
                item['kwic_sentence'] = item['truku_sentence']
                item['positions'] = []
        results.append(item)
        
    collocations = []
    ngrams = {'bigrams': [], 'trigrams': []}
    positions = {'initial': 0, 'medial': 0, 'final': 0}
    morphology = None
    
    sentences_for_stats = [r['truku_sentence'] for r in exact_rows if r['truku_sentence']]
    
    if search_type == 'truku':
        if is_proximity:
            collocations = get_collocations_prox(sentences_for_stats, word1, word2)
            positions = {}
        else:
            collocations = get_collocations_regex(sentences_for_stats, regex_exact)
            ngrams['bigrams'] = get_ngrams_regex(sentences_for_stats, regex_exact, 2)
            ngrams['trigrams'] = get_ngrams_regex(sentences_for_stats, regex_exact, 3)
            positions = get_position_stats_regex(sentences_for_stats, regex_exact)
            
            clean_word = re.sub(r'[^\w\-]', '', query).lower()
            if clean_word and match_mode == 'exact' and '*' not in clean_word:
                morphology = get_morphology_for_word(clean_word)
    else:
        colloc_counter = Counter()
        for sent in sentences_for_stats:
            tokens = set(tokenize(sent))
            for t in tokens:
                if t and len(t) >= 2:
                    colloc_counter[t] += 1
        collocations = [w for w, c in colloc_counter.most_common(5)]
        
    corpus_breakdown = {}
    for r in final_rows:
        src = r['corpus_source']
        if src not in corpus_breakdown:
            corpus_breakdown[src] = {'total': 0, 'exact': 0, 'partial': 0, 'subcategories': {}}
        corpus_breakdown[src]['total'] += 1
        
        if src in ['族語E樂園', '語推組織語料'] and r['detailed_source']:
            subcat = r['detailed_source'].replace('族語E樂園 - ', '').strip()
            if subcat:
                if subcat not in corpus_breakdown[src]['subcategories']:
                    corpus_breakdown[src]['subcategories'][subcat] = 0
                corpus_breakdown[src]['subcategories'][subcat] += 1
        
    for r in exact_rows:
        src = r['corpus_source']
        if src in corpus_breakdown: corpus_breakdown[src]['exact'] += 1
        
    for r in partial_rows:
        src = r['corpus_source']
        if src in corpus_breakdown: corpus_breakdown[src]['partial'] += 1
        
    return {
        "results": results,
        "total_occurrences": len(final_rows),
        "exact_occurrences": len(exact_rows),
        "partial_occurrences": len(partial_rows),
        "corpus_breakdown": corpus_breakdown,
        "collocations": collocations,
        "ngrams": ngrams,
        "positions": positions,
        "morphology": morphology
    }

@app.route('/')
def index():
    return app.send_static_file('index.html')

@app.route('/api/ai_analysis', methods=['GET', 'POST'])
def ai_analysis():
    custom_api_key = request.args.get('custom_api_key', '')
    
    keys_pool = []
    if custom_api_key:
        keys_pool.append(custom_api_key)
        
    env_keys_str = os.environ.get("GEMINI_API_KEY", "")
    if env_keys_str and env_keys_str != '請在此填寫您的API_KEY':
        for k in env_keys_str.split(','):
            if k.strip(): keys_pool.append(k.strip())
            
    for k, v in os.environ.items():
        if k.startswith("GEMINI_API_KEY_") and v.strip():
            for key in v.split(','):
                if key.strip(): keys_pool.append(key.strip())
                
    if not keys_pool:
        return jsonify({"error": "尚未設定 Gemini API Key。請在 .env 檔案中設定您的 GEMINI_API_KEY。"}), 400
        
    try:
        from google import genai
        from google.genai import types
        
        word1 = request.args.get('w1', '').strip()
        word2 = request.args.get('w2', '').strip()
        
        try:
            limit = int(request.args.get('limit', 20))
        except ValueError:
            limit = 20
        
        if not word1:
            return jsonify({"error": "請至少提供一個詞彙進行分析"}), 400
            
        conn = get_db_connection()
        cursor = conn.cursor()
        
        def fetch_sentences(word):
            safe_q = re.escape(word)
            regex_exact = re.compile(r'\b' + safe_q + r'\b', re.IGNORECASE)
            sql = "SELECT truku_sentence, chinese_translation FROM dictionary WHERE truku_sentence LIKE ?"
            all_rows = cursor.execute(sql, (f"%{word}%",)).fetchall()
            
            sentences = []
            for r in all_rows:
                if r['truku_sentence'] and regex_exact.search(r['truku_sentence']):
                    sentences.append(f"- 族語：{r['truku_sentence']}\n  中文：{r['chinese_translation']}")
                    if len(sentences) >= limit:
                        break
            return sentences
            
        w1_sentences = fetch_sentences(word1)
        w2_sentences = fetch_sentences(word2) if word2 else []
        
        conn.close()
        
        if not w1_sentences:
            return jsonify({"error": f"找不到關於「{word1}」的例句，無法進行分析。"}), 404
            
        if word2 and not w2_sentences:
            return jsonify({"error": f"找不到關於「{word2}」的例句，無法進行分析。"}), 404
            
        prompt = (
            "你是一位極度專業的「太魯閣語 (Truku) 語言學家」，擅長透過語料庫進行語法與語義分析。\n"
            "請直接輸出分析報告，不要重複或覆述我的指令。\n"
            "請根據以下從實際語料庫抽取的例句，進行極度詳細、深度且具體的分析。你的分析必須完全基於提供的例句，不可憑空捏造。\n\n"
        )
        
        if word2:
            prompt += f"【分析目標】：比較「{word1}」與「{word2}」在語意、語法與使用情境上的差異。\n\n"
            prompt += f"【{word1} 的實際例句 (共 {len(w1_sentences)} 句)】\n" + "\n".join(w1_sentences) + "\n\n"
            prompt += f"【{word2} 的實際例句 (共 {len(w2_sentences)} 句)】\n" + "\n".join(w2_sentences) + "\n\n"
            prompt += (
                "【輸出格式要求】\n"
                "請使用清晰的 Markdown 格式，並包含以下架構：\n"
                "### 1. 核心語意差異 (Core Semantic Differences)\n請說明兩者在最根本的意思與使用情境上有何不同。\n"
                "### 2. 搭配詞與語法分析 (Collocation & Syntactic Analysis)\n請分析這兩個詞通常和什麼樣的動詞、名詞或介系詞搭配使用。\n"
                "### 3. 實際例句對照解析 (Comparative Examples)\n請「務必」從我提供的例句中，各挑選 2-3 句最經典的例句（包含族語與中文），具體說明為什麼這裡只能用 A 而不能用 B，或是兩者換用的話語氣會有何不同。\n"
                "### 4. 總結建議 (Summary)\n給學習者的一句話總結，教他們如何快速區分這兩個詞。\n"
            )
        else:
            prompt += f"【分析目標】：分析太魯閣語單字「{word1}」的語意、語法特色與常見使用情境。\n\n"
            prompt += f"【{word1} 的實際例句 (共 {len(w1_sentences)} 句)】\n" + "\n".join(w1_sentences) + "\n\n"
            prompt += (
                "【輸出格式要求】\n"
                "請使用清晰的 Markdown 格式，並包含以下架構：\n"
                "### 1. 核心語意 (Core Meaning)\n請精準總結這個詞的根本意思與使用情境。\n"
                "### 2. 搭配詞與語法特性 (Collocations & Grammar)\n請分析這個詞通常和什麼樣的詞彙搭配？在句子中扮演什麼角色？\n"
                "### 3. 經典例句解析 (Key Examples Analysis)\n請「務必」從我提供的例句中，挑選 3-4 句最經典的例句（包含族語與中文），並進行深度的剖析，解釋其在不同上下文的語氣變化。\n"
            )
            
        system_instruction = "你是一位精通太魯閣語與繁體中文的專業語言學家。請直接輸出繁體中文的語言學分析報告，不要包含任何思考過程或確認指令的字眼。"
        
        last_exception = None
        response = None
        
        for current_key in keys_pool:
            try:
                client = genai.Client(api_key=current_key)
                response = client.models.generate_content(
                    model='gemini-3.5-flash',
                    contents=prompt,
                    config=types.GenerateContentConfig(
                        system_instruction=system_instruction,
                        temperature=0.2,
                        top_p=0.95,
                        top_k=40,
                        max_output_tokens=8192
                    )
                )
                break
            except Exception as e:
                print(f"⚠️ API Key Error ({current_key[:5]}...): {e}")
                last_exception = e
                continue
                
        if not response:
            raise last_exception
        
        return jsonify({
            "result": response.text,
            "w1_count": len(w1_sentences),
            "w2_count": len(w2_sentences)
        })
        
    except Exception as e:
        return jsonify({"error": f"AI 分析發生錯誤: {str(e)}"}), 500

@app.route('/api/stats', methods=['GET'])
def get_stats():
    global stats_cache
    if stats_cache is not None: return jsonify(stats_cache)
    try:
        return _compute_stats()
    except Exception as e:
        import traceback
        traceback.print_exc()
        return jsonify({"error": f"統計資料計算失敗: {str(e)}"}), 500

def _compute_stats():
    global stats_cache
    conn = get_db_connection()
    cursor = conn.cursor()

    total_entries = cursor.execute("SELECT count(*) FROM dictionary").fetchone()[0]
    total_sentences = cursor.execute("SELECT count(*) FROM dictionary WHERE truku_sentence IS NOT NULL AND truku_sentence != ''").fetchone()[0]

    c_col = get_corpus_col(cursor)
    if c_col:
        corpus_dist_raw = cursor.execute(f"SELECT {c_col}, count(*) FROM dictionary GROUP BY {c_col}").fetchall()
        corpus_distribution = {row[0]: row[1] for row in corpus_dist_raw}
    else:
        corpus_distribution = {"太魯閣語辭典": total_entries}
    
    if has_column(cursor, 'dictionary', 'source_word'):
        source_words = [row[0] for row in cursor.execute("SELECT source_word FROM dictionary WHERE source_word IS NOT NULL AND source_word != ''").fetchall()]
    else:
        source_words = []
    first_letter_counter = Counter()
    prefix_counter = {'m-': 0, 'p-': 0, 'emp-': 0, 't-': 0}
    suffix_counter = {'-an': 0, '-ay': 0, '-i': 0, '-un': 0, '-aw': 0, '-a': 0}
    
    for w in source_words:
        clean = re.sub(r'[^\w]', '', w).lower()
        if clean: 
            first_letter_counter[clean[0]] += 1
            
            if clean.startswith('emp'): prefix_counter['emp-'] += 1
            elif clean.startswith('m'): prefix_counter['m-'] += 1
            elif clean.startswith('p'): prefix_counter['p-'] += 1
            elif clean.startswith('t'): prefix_counter['t-'] += 1
            
            if clean.endswith('an'): suffix_counter['-an'] += 1
            elif clean.endswith('ay'): suffix_counter['-ay'] += 1
            elif clean.endswith('un'): suffix_counter['-un'] += 1
            elif clean.endswith('aw'): suffix_counter['-aw'] += 1
            elif clean.endswith('a'): suffix_counter['-a'] += 1
            elif clean.endswith('i'): suffix_counter['-i'] += 1
            
    c_col = get_corpus_col(cursor)
    if c_col:
        sentences_raw = cursor.execute(f"SELECT truku_sentence, {c_col} as corpus_source FROM dictionary WHERE truku_sentence IS NOT NULL AND truku_sentence != ''").fetchall()
    else:
        sentences_raw = [(row[0], '太魯閣語辭典') for row in cursor.execute("SELECT truku_sentence FROM dictionary WHERE truku_sentence IS NOT NULL AND truku_sentence != ''").fetchall()]
    
    if has_column(cursor, 'dictionary', 'source_word_translation'):
        translations_raw = cursor.execute("SELECT source_word, source_word_translation FROM dictionary WHERE source_word IS NOT NULL AND source_word != ''").fetchall()
    else:
        translations_raw = []
    conn.close()
    
    translation_map = {}
    for w, t in translations_raw:
        if w and t:
            clean_t = t.split('；')[0].split(';')[0].strip()
            translation_map[w.lower()] = clean_t
            
    word_freq_counter = Counter()
    bigram_counter = Counter()
    trigram_counter = Counter()
    total_word_count = 0
    
    corpus_word_counters = {
        '族語辭典': Counter(),
        '族語E樂園': Counter(),
        '族語文學': Counter()
    }
    corpus_total_words = {
        '族語辭典': 0,
        '族語E樂園': 0,
        '族語文學': 0
    }
    
    for sent, src in sentences_raw:
        tokens = tokenize(sent)
        total_word_count += len(tokens)
        
        if src in corpus_total_words:
            corpus_total_words[src] += len(tokens)
            
        for t in tokens:
            if len(t) >= 2: 
                word_freq_counter[t] += 1
                if src in corpus_word_counters:
                    corpus_word_counters[src][t] += 1
            
        for i in range(len(tokens)):
            if i < len(tokens) - 1 and len(tokens[i]) >= 2 and len(tokens[i+1]) >= 2:
                bigram_counter[f"{tokens[i]} {tokens[i+1]}"] += 1
            if i < len(tokens) - 2 and len(tokens[i]) >= 2 and len(tokens[i+1]) >= 2 and len(tokens[i+2]) >= 2:
                trigram_counter[f"{tokens[i]} {tokens[i+1]} {tokens[i+2]}"] += 1
                
    avg_sentence_length = round(total_word_count / total_sentences, 1) if total_sentences > 0 else 0
    
    top_words_raw = word_freq_counter.most_common(50)
    top_words = []
    for w, count in top_words_raw:
        trans = translation_map.get(w.lower(), "")
        top_words.append({"word": w, "count": count, "translation": trans})
        
    top_words_by_corpus = {}
    for corpus, counter in corpus_word_counters.items():
        top_words_by_corpus[corpus] = []
        total_words = corpus_total_words.get(corpus, 1)
        if total_words == 0: total_words = 1
        for w, count in counter.most_common(50):
            trans = translation_map.get(w.lower(), "")
            percentage = round((count / total_words) * 100, 2)
            top_words_by_corpus[corpus].append({"word": w, "count": count, "translation": trans, "percentage": percentage})
        
    stats_cache = {
        'total_entries': total_entries,
        'total_sentences': total_sentences,
        'avg_sentence_length': avg_sentence_length,
        'corpus_distribution': corpus_distribution,
        'first_letters': dict(first_letter_counter.most_common()),
        'prefixes': prefix_counter,
        'suffixes': suffix_counter,
        'top_words': top_words,
        'top_words_by_corpus': top_words_by_corpus,
        'top_bigrams': dict(bigram_counter.most_common(10)),
        'top_trigrams': dict(trigram_counter.most_common(10))
    }
    return jsonify(stats_cache)

@app.route('/api/v1/search', methods=['GET'])
@app.route('/api/search', methods=['GET'])
def search():
    query = request.args.get('q', '').strip()
    search_type = request.args.get('type', 'truku')
    match_mode = request.args.get('match_mode', 'exact')
    corpora = request.args.getlist('corpora[]')
    if not corpora:
        corpora_str = request.args.get('corpora')
        if corpora_str:
            corpora = corpora_str.split(',')
    data = perform_search(query, search_type, match_mode, limit=3000, corpora=corpora)
    return jsonify(data)

@app.route('/api/report', methods=['POST'])
def submit_report():
    data = request.json
    try:
        conn = get_db_connection()
        cursor = conn.cursor()
        cursor.execute('''
            CREATE TABLE IF NOT EXISTS feedback (
                id INTEGER PRIMARY KEY AUTOINCREMENT,
                item_id INTEGER,
                type TEXT NOT NULL,
                truku_sentence TEXT,
                chinese_translation TEXT,
                description TEXT NOT NULL,
                email TEXT,
                created_at TIMESTAMP DEFAULT CURRENT_TIMESTAMP
            )
        ''')
        cursor.execute('''
            INSERT INTO feedback (type, description, email, sentence_id, sentence_text)
            VALUES (?, ?, ?, ?, ?)
        ''', (data.get('type'), data.get('description'), data.get('email'), data.get('sentenceId'), data.get('sentence')))
        conn.commit()
        conn.close()
        
        # Here we could use smtplib to actually send the email to lowking@ilrdf.org.tw
        # Since we don't have SMTP credentials, saving to DB is the primary persistence.
        app.logger.info(f"Feedback received: {data}")
        
        return jsonify({"status": "success"})
    except Exception as e:
        app.logger.error(f"Feedback error: {str(e)}")
        return jsonify({"error": str(e)}), 500

@app.route('/api/export', methods=['GET'])
def export_csv():
    query = request.args.get('q', '').strip()
    search_type = request.args.get('type', 'truku')
    match_mode = request.args.get('match_mode', 'exact')
    corpora = request.args.getlist('corpora[]')
    if not corpora:
        corpora_str = request.args.get('corpora')
        if corpora_str:
            corpora = corpora_str.split(',')
    data = perform_search(query, search_type, match_mode, limit=None, corpora=corpora)
    si = StringIO()
    si.write('\ufeff')
    cw = csv.writer(si)
    cw.writerow(['ID', '匹配模式', '語料來源', '族語例句', '中文翻譯', '來源詞條', '詞彙位置', '來源檔案'])
    for item in data['results']:
        match_type = '精準 (Exact)' if item.get('is_exact') else '部分 (Partial)'
        positions = ', '.join(item.get('positions', []))
        cw.writerow([
            item.get('id', ''), match_type, item.get('corpus_source', ''), item.get('truku_sentence', ''),
            item.get('chinese_translation', ''), item.get('source_word', ''),
            positions, item.get('detailed_source', '')
        ])
    output = si.getvalue()
    si.close()
    return Response(
        output, mimetype="text/csv",
        headers={"Content-disposition": f"attachment; filename=truku_search_export.csv"}
    )

@app.route('/api/dictionary/add_sentence', methods=['POST'])
def api_dict_add_sentence():
    data = request.json
    truku_sentence = data.get('truku_sentence')
    chinese_translation = data.get('chinese_translation')
    corpus_source = data.get('corpus_source', '語料標注平台')
    
    if not truku_sentence or not chinese_translation:
        return jsonify({"error": "Missing sentence or translation"}), 400
        
    try:
        conn = sqlite3.connect(DB_FILE)
        c = conn.cursor()
        c.execute('''
            INSERT INTO dictionary (truku_sentence, chinese_translation, corpus_source)
            VALUES (?, ?, ?)
        ''', (truku_sentence, chinese_translation, corpus_source))
        conn.commit()
        conn.close()
        return jsonify({"status": "success"})
    except Exception as e:
        return jsonify({"error": str(e)}), 500

@app.route('/api/dictionary/add_word', methods=['POST'])
def api_dict_add_word():
    data = request.json
    source_word = data.get('source_word')
    source_word_translation = data.get('source_word_translation')
    corpus_source = data.get('corpus_source', '語料標注平台')
    
    if not source_word or not source_word_translation:
        return jsonify({"error": "Missing word or translation"}), 400
        
    try:
        conn = sqlite3.connect(DB_FILE)
        c = conn.cursor()
        c.execute('SELECT id FROM dictionary WHERE source_word = ?', (source_word,))
        if c.fetchone():
            conn.close()
            return jsonify({"status": "exists", "message": "Word already in dictionary"})
            
        c.execute('''
            INSERT INTO dictionary (truku_sentence, chinese_translation, source_word, source_word_translation, corpus_source)
            VALUES (?, ?, ?, ?, ?)
        ''', ('', '', source_word, source_word_translation, corpus_source))
        conn.commit()
        conn.close()
        return jsonify({"status": "success", "message": "Word added to dictionary"})
    except Exception as e:
        return jsonify({"error": str(e)}), 500

@app.route('/api/dictionary/check_sentence', methods=['POST'])
def api_dict_check_sentence():
    data = request.json
    truku_sentence = data.get('truku_sentence')
    if not truku_sentence:
        return jsonify({"exists": True})
        
    try:
        conn = sqlite3.connect(DB_FILE)
        c = conn.cursor()
        c.execute('SELECT id FROM dictionary WHERE truku_sentence = ?', (truku_sentence,))
        row = c.fetchone()
        conn.close()
        return jsonify({"exists": row is not None})
    except Exception as e:
        return jsonify({"error": str(e)}), 500

# --- Integrated AI Glosser Endpoints (Ported from main.py) ---
import time
import io
import json
from docx import Document
from docx.shared import Pt, Cm
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from flask import send_file

MEMORY_DB_FILE = "truku_memory.db"

def init_memory_db():
    conn = sqlite3.connect(MEMORY_DB_FILE)
    cursor = conn.cursor()
    cursor.execute('''
        CREATE TABLE IF NOT EXISTS verified_words (
            raw_word TEXT PRIMARY KEY,
            verified_morph TEXT,
            verified_gloss TEXT
        )
    ''')
    cursor.execute("CREATE INDEX IF NOT EXISTS idx_raw_word ON verified_words(raw_word)")
    conn.commit()
    conn.close()

init_memory_db()

BASE_SYSTEM_INSTRUCTION = (
    "你是一位精通台灣原住民族語的語言學權威專家，專精太魯閣語（Truku）。\n"
    "【核心語法指南】：\n"
    "1. 焦點系統：主事焦點（m- / <m> / em-）、受事焦點（-un ）、處所焦點（-an）、工具/受惠焦點（s-）。\n"
    "2. 時態與語體：進行狀態（gaga, ga, nii, gisu）、完成（-n- / wada）、未來（emp-）。\n"
    "3. 代名詞附著：主格/屬格（=ku 我.主格, =mu 我.屬格, =su 你.主格, =na 他/她.屬格, =dha 他們.屬格）。\n"
    "4. 格位標記：主格（ka）、主題標記（o）。\n\n"
    "【標記語言限制 - 全中文】：\n"
    "你的 gloss（語法標記）欄位「絕對不允許」出現任何英文簡寫。必須全部轉換為「中文語法術語」與「中文詞義」。\n"
    "複合語義一般用點號（.）連結（如 進行.助動）。\n"
    "❗ 極度重要：詞綴標記符號必須與第二行切分完全對應：\n"
    "  (1) 若是前綴或後綴 (如 m- 或 -un)，第三行的語法標記必須使用「連字號 (-)」連結。例如：m-niq 標記為 主事焦點-住，n-apa 標記為 完成貌-背。\n"
    "  (2) 若是中綴 (如 <m> 或 <n>)，第三行的語法標記必須使用「角括號 (< >)」將功能包起來。例如：d<m>udug 標記為 <主事焦點>推動，s<n>alu 標記為 <完成貌>製作。\n"
    "代名詞附著請加等號（如 =su）。\n"
    "【輸出格式規範】：必須嚴格遵守 JSON Schema 返回 sentences 陣列。"
)

@app.route('/api/parse', methods=['POST'])
def parse_corpus():
    data = request.json or {}
    truku_sentence = data.get('truku_sentence', '')
    chinese_translation = data.get('chinese_translation', '')
    custom_api_key = data.get('custom_api_key', '')
    
    keys_pool = []
    if custom_api_key:
        keys_pool.append(custom_api_key)
        
    env_keys_str = os.environ.get("GEMINI_API_KEY", "")
    if env_keys_str:
        for k in env_keys_str.split(','):
            if k.strip(): keys_pool.append(k.strip())
            
    for k, v in os.environ.items():
        if k.startswith("GEMINI_API_KEY_") and v.strip():
            for key in v.split(','):
                if key.strip(): keys_pool.append(key.strip())
                
    if not keys_pool:
        return jsonify({"detail": "未偵測到 GEMINI_API_KEY，且未提供自訂 Key"}), 500

    try:
        start_time = time.perf_counter()
        import string
        clean_sentence = truku_sentence.translate(str.maketrans('', '', string.punctuation))
        words_in_sentence = list(set(clean_sentence.lower().split()))
        
        # 1. 查詢記憶庫
        conn_mem = sqlite3.connect(MEMORY_DB_FILE)
        cursor_mem = conn_mem.cursor()
        memory_prompt_segment = "【黃金記憶庫校正範例 (最高優先級)】:\n"
        has_mem_data = False
        for word in words_in_sentence:
            cursor_mem.execute("SELECT verified_morph, verified_gloss FROM verified_words WHERE raw_word = ?", (word,))
            row = cursor_mem.fetchone()
            if row:
                memory_prompt_segment += f"- \"{word}\" -> 詞切分: \"{row[0]}\", 標記: \"{row[1]}\"\n"
                has_mem_data = True
        conn_mem.close()

        # 2. 查詢主辭典
        dict_prompt_segment = "【單字字典解釋參考】:\n"
        has_dict_data = False
        if os.path.exists(DB_FILE):
            conn_dict = sqlite3.connect(DB_FILE)
            cursor_dict = conn_dict.cursor()
            if has_column(cursor_dict, 'dictionary', 'source_word_translation'):
                for word in words_in_sentence:
                    cursor_dict.execute("SELECT source_word_translation FROM dictionary WHERE source_word = ? LIMIT 1", (word,))
                    row = cursor_dict.fetchone()
                    if row and row[0]:
                        dict_prompt_segment += f"- \"{word}\" 的字典解釋: \"{row[0]}\"\n"
                        has_dict_data = True
            conn_dict.close()

        # 組合 Prompt
        user_prompt = f"請精確標註以下語料：\n{truku_sentence}\n\n"
        if chinese_translation:
            user_prompt += f"參考意譯（請務必依據此意譯進行標註）：\n{chinese_translation}\n\n"
        
        if has_dict_data:
            user_prompt = f"{dict_prompt_segment}\n{user_prompt}"
        if has_mem_data:
            user_prompt = f"{memory_prompt_segment}\n{user_prompt}"

        from google import genai
        from google.genai import types
        
        package_schema = {
            "type": "OBJECT",
            "properties": {
                "sentences": {
                    "type": "ARRAY",
                    "items": {
                        "type": "OBJECT",
                        "properties": {
                            "words": {
                                "type": "ARRAY",
                                "items": {
                                    "type": "OBJECT",
                                    "properties": {
                                        "raw": {"type": "STRING"},
                                        "morph": {"type": "STRING"},
                                        "gloss": {"type": "STRING"},
                                        "confidence": {"type": "STRING"}
                                    },
                                    "required": ["raw", "morph", "gloss", "confidence"]
                                }
                            },
                            "translation": {"type": "STRING"}
                        },
                        "required": ["words", "translation"]
                    }
                }
            },
            "required": ["sentences"]
        }

        last_exception = None
        response = None
        
        for current_key in keys_pool:
            try:
                client = genai.Client(api_key=current_key)
                response = client.models.generate_content(
                    model='gemini-3.5-flash',
                    contents=user_prompt,
                    config=types.GenerateContentConfig(
                        system_instruction=BASE_SYSTEM_INSTRUCTION,
                        response_mime_type="application/json",
                        response_schema=package_schema, 
                        temperature=0.0
                    )
                )
                break
            except Exception as e:
                print(f"⚠️ API Key Error ({current_key[:5]}...): {e}")
                last_exception = e
                continue
                
        if not response:
            raise last_exception

        parsed_json = json.loads(response.text.strip())
        
        if chinese_translation:
            import re
            chinese_sentences = [s.strip() for s in re.split(r'(?<=[。！？!?\n])\s*', chinese_translation) if s.strip()]
            sentences = parsed_json.get("sentences", [])
            for i, sentence in enumerate(sentences):
                if i < len(chinese_sentences):
                    sentence["translation"] = chinese_sentences[i]
                
        parsed_json["elapsed_time"] = round(time.perf_counter() - start_time, 2)
        return jsonify(parsed_json)

    except Exception as e:
        error_str = str(e).lower()
        print(f"❌ 分析錯誤: {e}")
        if "429" in error_str or "quota" in error_str or "exhausted" in error_str:
            return jsonify({"detail": "QUOTA_EXCEEDED"}), 429
        return jsonify({"detail": str(e)}), 500

@app.route('/api/save_correct', methods=['POST'])
def save_correct():
    try:
        data = request.json or {}
        words = data.get('words', [])
        conn = sqlite3.connect(MEMORY_DB_FILE)
        cursor = conn.cursor()
        for word in words:
            import string
            raw = word.get('raw', '')
            clean_raw = raw.translate(str.maketrans('', '', string.punctuation)).lower()
            if not clean_raw: continue
            cursor.execute('''
                INSERT OR REPLACE INTO verified_words (raw_word, verified_morph, verified_gloss)
                VALUES (?, ?, ?)
            ''', (clean_raw, word.get('morph', ''), word.get('gloss', '')))
        conn.commit()
        conn.close()
        return jsonify({"status": "success"})
    except Exception as e:
        return jsonify({"detail": str(e)}), 500

@app.route('/api/download_word', methods=['POST'])
def download_word():
    try:
        data = request.json or {}
        sentences = data.get('sentences', [])
        doc = Document()
        style = doc.styles['Normal']
        style.font.name = 'Times New Roman'
        style.font.size = Pt(12)
        style.paragraph_format.space_after = Pt(0)
        style.paragraph_format.space_before = Pt(0)
        rPr = style.element.get_or_add_rPr()
        rFonts = OxmlElement('w:rFonts')
        rFonts.set(qn('w:ascii'), 'Times New Roman')
        rFonts.set(qn('w:hAnsi'), 'Times New Roman')
        rFonts.set(qn('w:eastAsia'), '標楷體')
        rPr.append(rFonts)

        section = doc.sections[0]
        section.left_margin = Cm(2.0)
        section.right_margin = Cm(2.0)

        for sentence in sentences:
            words = sentence.get('words', [])
            if len(words) == 0: continue
            
            raw_line = " ".join([str(w.get('raw','')) for w in words if w.get('raw')]).strip()
            if not raw_line.endswith('.'): raw_line += '.'
            doc.add_paragraph(raw_line)
            
            merged_cols = []
            for w in words:
                m = str(w.get('morph', '')).strip()
                g = str(w.get('gloss', '')).strip()
                if m.startswith('=') and merged_cols:
                    merged_cols[-1]['morph'] += m
                    merged_cols[-1]['gloss'] += g
                else:
                    merged_cols.append({'morph': m, 'gloss': g})
            
            chunk_size = 6
            for chunk_start in range(0, len(merged_cols), chunk_size):
                chunk_cols = merged_cols[chunk_start:chunk_start+chunk_size]
                table = doc.add_table(rows=2, cols=len(chunk_cols))
                table.autofit = True
                tblW = table._element.tblPr.xpath('w:tblW')
                if tblW:
                    tblW[0].set(qn('w:type'), 'auto')
                    tblW[0].set(qn('w:w'), '0')
                for row in table.rows:
                    for cell in row.cells:
                        tcW = cell._element.get_or_add_tcPr().xpath('w:tcW')
                        if tcW:
                            tcW[0].set(qn('w:type'), 'auto')
                            tcW[0].set(qn('w:w'), '0')

                for i, col in enumerate(chunk_cols):
                    c0 = table.cell(0, i)
                    c0.text = col['morph']
                    c0._element.get_or_add_tcPr().append(OxmlElement('w:noWrap'))
                    c1 = table.cell(1, i)
                    c1.text = col['gloss']
                    c1._element.get_or_add_tcPr().append(OxmlElement('w:noWrap'))
                
            doc.add_paragraph(str(sentence.get('translation', '')))
            doc.add_paragraph()

        bio = io.BytesIO()
        doc.save(bio)
        bio.seek(0)
        return send_file(
            bio,
            mimetype="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            as_attachment=True,
            download_name="glossed_output.docx"
        )
    except Exception as e:
        return jsonify({"detail": str(e)}), 500

if __name__ == '__main__':
    os.makedirs('static', exist_ok=True)
    build_morphology_cache() # Preheat cache
    app.run(debug=True, port=5000)
